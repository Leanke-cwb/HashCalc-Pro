from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from datetime import datetime
from pathlib import Path
import os


class CertidaoHash:


    @staticmethod
    def gerar(arquivo, hashes):


        caminho = Path(arquivo)


        nome_saida = caminho.with_name(
            f"Certidao_Hash_{caminho.stem}.docx"
        )


        documento = Document()



        # Título

        titulo = documento.add_heading(
            "CERTIDÃO DE HASH DIGITAL",
            level=1
        )


        titulo.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )



        documento.add_paragraph(
            ""
        )



        texto = (
            "CERTIFICO, para os devidos fins, "
            "que foi realizado o cálculo de "
            "integridade criptográfica do arquivo "
            "abaixo descrito."
        )


        documento.add_paragraph(
            texto
        )



        documento.add_heading(
            "1. IDENTIFICAÇÃO DO ARQUIVO",
            level=2
        )


        documento.add_paragraph(
            f"Nome do arquivo: {caminho.name}"
        )


        documento.add_paragraph(
            f"Caminho: {arquivo}"
        )



        tamanho = os.path.getsize(
            arquivo
        )


        documento.add_paragraph(
            f"Tamanho: {tamanho:,} bytes"
        )



        documento.add_paragraph(
            "Data do cálculo: "
            +
            datetime.now().strftime(
                "%d/%m/%Y %H:%M:%S"
            )
        )



        documento.add_heading(
            "2. ALGORITMOS UTILIZADOS",
            level=2
        )


        for algoritmo in hashes.keys():

            documento.add_paragraph(
                algoritmo.upper()
            )



        documento.add_heading(
            "3. VALORES HASH",
            level=2
        )


        for algoritmo, valor in hashes.items():


            documento.add_paragraph(
                algoritmo.upper()
            )


            paragrafo = documento.add_paragraph(
                valor
            )


            for run in paragrafo.runs:

                run.font.name = "Courier New"
                run.font.size = Pt(9)



        documento.add_heading(
            "4. DECLARAÇÃO",
            level=2
        )


        documento.add_paragraph(

            "Declaro que os valores hash apresentados "
            "foram obtidos através de algoritmos "
            "criptográficos utilizados para verificação "
            "da integridade do arquivo analisado."

        )



        documento.add_paragraph(
            ""
        )


        documento.add_paragraph(
            "__________________________________"
        )


        documento.add_paragraph(
            "Responsável pelo cálculo"
        )



        documento.add_paragraph(
            ""
        )


        documento.add_paragraph(
            "Documento gerado pelo HashCalc Pro"
        )



        documento.save(
            nome_saida
        )



        return str(nome_saida)